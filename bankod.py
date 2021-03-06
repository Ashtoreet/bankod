import requests
import time
import dateparser
from docx import Document
import cyrtranslit
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException

from bs4 import BeautifulSoup as bs

from datetime import datetime as dt
from fake_useragent import UserAgent

from operator import itemgetter


# from news import news


def s_get_url(bank_name):
    driver = webdriver.Chrome()
    driver.get("https://www.bankodrom.ru")
    time.sleep(1)
    c = driver.find_element_by_id('term')
    for i in bank_name:
        time.sleep(.1)
        c.send_keys(i)

    time.sleep(3)

    try:
        driver.find_element_by_id('onesignal-popover-cancel-button').click()
    except NoSuchElementException:
        print()

    time.sleep(1)

    try:
        f = driver.find_element_by_class_name('bank-1').click()
        url = driver.current_url
        print(url)
    except Exception:
        driver.find_element_by_name('button.onesignal-popover-cancel-button').click()
        f
        url
        print('url', url)

    time.sleep(1)
    # print(input('close? ', ))
    driver.close()

    return url


def user_a():
    ua = UserAgent()
    header = {'User-Agent':str(ua.chrome)}
    return header


def get_news(url, page_number, header):
    r = requests.get(url + '{}{}'.format('novosti/?p=', page_number), headers=header).text
    soup = bs(r, 'lxml')

    news = soup.find('div', class_='bank-news-items')
    if news:
        print('получили блок news')
        # print(news)
        return news
    else:
        print('news', news)


def last_date_on_page(news):
    last_date = news.findAll('td', 'date-title').pop()
    return dateparser.parse(last_date.text)


def find_not_all_news(start_date, news):
    for i in news.findAll('td'):
        if i.attrs.get('class') == ['date-title'] and dateparser.parse(i.text) < start_date:
            print('Дата предыдущих новостей: ', i.text)
            articles = i.find_all_previous('td', class_="news-title")
            if articles:
                print('список статей: ', len(articles))
                articles.reverse()
                return articles
            else:
                print()
                # print('articles', articles)


def find_all_news(news):
    articles = news.find_all('td', class_="news-title")
    if articles:
        print('список статей: ', len(articles))
    # articles.reverse()
        return articles
    else:
        print()
        # print('articles: ', articles)


def page_navigation(start_date, bank_name, url):
    print('page_navigation')
    # try:
    links = []
    page_number = 1
    while True:
        header = user_a()
        news = get_news(url, page_number, header)

        last_d_o_p = last_date_on_page(news)
        print('last_d_o_p' , last_d_o_p)

        if last_d_o_p > start_date:
            articles = find_all_news(news)
            for i in articles:
                if i.text.split()[-1] == 'Banki.ru':
                # print(i.text.split()[-1], i.text.split()[-1] == 'Banki.ru')
                    link = i.find('a').get('href')
                    links.append('{}{}'.format('https://www.bankodrom.ru', link))
                    print(i.text.split()[-1])
            page_number += 1
        else:
            articles = find_not_all_news(start_date, news)
            for i in articles:
                if i.text.split()[-1] == 'Banki.ru':
                # print(i.text.split()[-1], i.text.split()[-1] == 'Banki.ru')
                    print(i.text.split()[-1])
                    link = i.find('a').get('href')
                    links.append('{}{}'.format('https://www.bankodrom.ru', link))
            break
    print('links', links)
    return links


def array_list(links):
    parse_text = []
    parse_text_prev = []
    for i in links:
        header = user_a()
        url = requests.get(i, headers=header).text
        soup = bs(url, 'lxml')

        date = dateparser.parse(soup.find('span', class_='news-date').text).date()
        article_date = date.strftime("%d.%m.%Y")
        head = soup.find('h1').text

        body = []
        all_body = soup.find('div', class_='news-item-body')
        for child in all_body.children:
            if child.name is 'p':
                body.append(child.text)

        parse_text_prev.append({
            'date': article_date,
            'head': head,
            'body': body
        })
    sorted(parse_text_prev, key=lambda x: x['date'], reverse=True)
    print('получаем массив статей')

    return parse_text_prev


def to_word(parse_text, file):
    document = Document()
    for i in parse_text:
        document.add_paragraph(i['date'])
        document.add_heading(i['head'], level=1)
        for n in i['body']:
            document.add_paragraph(n)
        document.add_paragraph()
    document.save(file)
    print('сохраняем документ')



if __name__ == '__main__':
    bank_name = input('Введите название: ', )
    # bank_name = 'Промсвязьбанк'
    file = (cyrtranslit.to_latin(bank_name) + '_' + dt.today().strftime("%d-%m-%Y") + '.docx')

    url = s_get_url(bank_name)
    start_date = dateparser.parse('01 января 2018 г.')
    print(start_date)

    links = page_navigation(start_date, bank_name, url)
    parse_text = array_list(links)
    to_word(parse_text, file)
